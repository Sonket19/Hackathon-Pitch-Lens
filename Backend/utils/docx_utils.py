# from docx import Document
# from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import tempfile
# from typing import Dict
# import logging
# from utils.gcs_utils import GCSManager

# logger = logging.getLogger(__name__)

# class MemoExporter:
#     def __init__(self):
#         self.gcs_manager = GCSManager()

#     async def create_memo_docx(self, deal_id: str, memo_text: str) -> str:
#         """Create DOCX memo and upload to GCS"""
#         try:
#             # Create new document
#             doc = Document()

#             # Add title
#             title = doc.add_heading('Investment Analysis Memo', 0)
#             title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#             # Add memo content
#             self._add_memo_content(doc, memo_text)

#             # Save to temporary file
#             with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
#                 doc.save(temp_file.name)
#                 temp_docx_path = temp_file.name

#             try:
#                 # Upload to GCS
#                 gcs_path = f"deals/{deal_id}/memo.docx"

#                 with open(temp_docx_path, 'rb') as docx_file:
#                     blob = self.gcs_manager.bucket.blob(gcs_path)
#                     blob.upload_from_file(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

#                 return f"gs://{self.gcs_manager.bucket.name}/{gcs_path}"

#             finally:
#                 # Clean up temp file
#                 import os
#                 os.unlink(temp_docx_path)

#         except Exception as e:
#             logger.error(f"DOCX creation error: {str(e)}")
#             raise

#     def _add_memo_content(self, doc: Document, memo_text: str):
#         """Add formatted memo content to document"""
#         try:
#             # Split memo into sections
#             sections = memo_text.split('')

#             for section in sections:
#                 if not section.strip():
#                     continue

#                 lines = section.split('')
#                 first_line = lines[0].strip()

#                 # Check if this is a heading (starts with number or contains keywords)
#                 if (first_line and 
#                     (first_line[0].isdigit() or 
#                      any(keyword in first_line.lower() for keyword in 
#                          ['executive', 'summary', 'founder', 'problem', 'opportunity', 
#                           'differentiator', 'team', 'market', 'risks', 'recommendation']))):
#                     # Add as heading
#                     doc.add_heading(first_line, level=1)

#                     # Add remaining content as paragraphs
#                     for line in lines[1:]:
#                         if line.strip():
#                             if line.strip().startswith('•') or line.strip().startswith('-'):
#                                 # Bullet point
#                                 p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
#                             else:
#                                 # Regular paragraph
#                                 doc.add_paragraph(line.strip())
#                 else:
#                     # Regular content
#                     for line in lines:
#                         if line.strip():
#                             if line.strip().startswith('•') or line.strip().startswith('-'):
#                                 # Bullet point
#                                 p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
#                             else:
#                                 # Regular paragraph
#                                 doc.add_paragraph(line.strip())

#                 # Add spacing between sections
#                 doc.add_paragraph()

#         except Exception as e:
#             logger.error(f"Content formatting error: {str(e)}")
#             # Fallback: add raw text
#             doc.add_paragraph(memo_text)


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import logging
from utils.gcs_utils import GCSManager

logger = logging.getLogger(__name__)

class MemoExporter:
    def __init__(self):
        self.gcs_manager = GCSManager()

    async def create_memo_docx(self, deal_id: str, memo_json: dict) -> str:
        """Create DOCX memo from JSON and upload to GCS"""
        try:
            print("memo_json : ",memo_json);
            doc = Document()

            # Add title
            title = doc.add_heading('Investment Analysis Memo', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add content from JSON
            self._add_json_content(doc, memo_json)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                doc.save(temp_file.name)
                temp_docx_path = temp_file.name
            print("add_json_content done");
            try:
                # Upload to GCS
                gcs_path = f"deals/{deal_id}/memo.docx"
                with open(temp_docx_path, 'rb') as docx_file:
                    blob = self.gcs_manager.bucket.blob(gcs_path)
                    blob.upload_from_file(
                        docx_file,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                return f"gs://{self.gcs_manager.bucket.name}/{gcs_path}"

            finally:
                import os
                os.unlink(temp_docx_path)

        except Exception as e:
            logger.error(f"DOCX creation error: {str(e)}")
            raise

    def _add_json_content(self, doc: Document, data, level=1):
        """Recursively add JSON content to the document"""
        if isinstance(data, dict):
            for key, value in data.items():
                # Format key as heading
                heading_text = key.replace("_", " ").title()
                doc.add_heading(heading_text, level=level)

                # Recurse for value
                self._add_json_content(doc, value, level=level+1)

        elif isinstance(data, list):
            for item in data:
                if isinstance(item, (dict, list)):
                    self._add_json_content(doc, item, level=level)
                else:
                    # Add list item as bullet point
                    doc.add_paragraph(str(item), style='List Bullet')

        else:
            # Add simple value as paragraph
            doc.add_paragraph(str(data))
